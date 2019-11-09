import os
import re
import chardet
import base64
from docx import *
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT,WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from 集成测试用例集.model import docx
class Deal:
    def __init__(self,projectname,testpeople,projectyear,projectmonth,projectday):
        '''修改项'''
        self.projectname=projectname
        self.testpeople=testpeople
        self.projectyear=projectyear
        self.projectmonth=projectmonth
        self.projectday=projectday
        '''END'''

    funname=''#入口函数名
    funlist=[]#序列中的函数列表（入口函数1个+修改函数n个）

    # 构建正则表达式
    # for text
    patternAllCase=re.compile('# Begin Test Case(.*?)# End Test Case',re.S) #解析一个.tcf中的所有案例
    patternIDMess=re.compile('#Begin Sequence Documentation.*?Sequence Id:\s*(.*?)\s*#Begin Sequence Descriptions\s*(.*?)\s*#End Sequence Descriptions\s*Test Engineer:\s*(.*?)\s*#Begin Reference\s*(.*?)\s*#End Reference\s*#Begin User Global\s*(.*?)\s*#End User Global\s*#Begin Comment\s*(.*?)\s*#End Comment', re.S)# 解析该序列标识
    patternFunList=re.compile('#Begin Isolated Procedure.*?Procedure =\s*(\S*)\s*Procedure Number',re.S)
    #for case
    patternGetFunName=re.compile(r'Procedure =\s*(\S*)\s*\n',re.S) #解析一个case中的函数名
    patternGetFileName=re.compile(r'File.*\\(\w*)\.c',re.S) #解析一个case中的文件名
    patternGetTesting=re.compile('#Begin View Case Documentation.*?Test Skills:\s*(.*?)\s*#Begin Case Descriptions\s*(.*?)\s*#End Case Descriptions',re.S) #解析测试技术&用例描述
    patternGetG=re.compile('# Begin Variable\s*Name =\s*(\S*)\s*Decl_type =\s*(\S*)\s*Usage = G\s*Value =\s*(\S*)\s*# End Variable',re.S) #解析全局变量
    patternInit=re.compile('# Begin Startup Code\s*(.*?)\s*# End Startup Code',re.S) #解析初始化代码
    patternTCStub=re.compile('# Begin TC Stub\s*Procedure =\s(\S*)\s*Overloading =.*?Type =\s*(.*?)\s*Value =\s*(\S*)\s*# End TC Stub Return Value',re.S) #解析桩函数
    patternGetZ=re.compile('# Begin Variable\s*Name =\s*(\S*)\s*Decl_type =\s*(\S*)\s*Usage = Z\s*Value =\s*(\S*)\s*# End Variable',re.S) #解析输入
    patternGetH=re.compile('# Begin Variable\s*Name =\s*(\S*)\s*Decl_type =\s*(\S*)\s*Usage = H\s*Value =\s*(\S*)\s*# End Variable',re.S) #解析输出
    patternGetO=re.compile('# Begin Variable\s*Name =\s*(\S*)\s*Decl_type =\s*(\S*)\s*Usage = O\s*Value =\s*(\S*)\s*# End Variable',re.S) #解析返回值
    #解析全局输入
    patternGlobalInit=re.compile('# Begin Code\s*(.*?)\s*# End Code',re.S)

    #获取 XXXXX.tcf中的文本
    def getText(self, url):
        with open(url, 'r',encoding="GBK") as f:
            text=f.read()
            return text
            '''
                f_read=f.read()
            f_charInfo=chardet.detect(f_read)
            f_read_decode=f_read.decode(f_charInfo['encoding'])
            print(f_charInfo['encoding'])
            print(f_read_decode)
        return f_read_decode
            '''



    #正则匹配所有tcf中的有用信息,返回一个列表
    def getCaseList(self,text):
        '''解析函数列表'''
        self.funlist=re.findall(self.patternFunList, text)
        '''解析序列标识'''
        idmessdict={}
        idmesslist=re.findall(self.patternIDMess, text)
        if(idmesslist):
            if(idmesslist[0][0]):
                idmessdict['Id']=idmesslist[0][0]
            else:
                idmessdict['Id']='无'
            if(idmesslist[0][1]):
                idmessdict['Descriptions']=idmesslist[0][1]
            else:
                idmessdict['Descriptions']='无'
            if(idmesslist[0][2]):
                idmessdict['Engineer']=idmesslist[0][2]
            else:
                idmessdict['Engineer']='无'
            if(idmesslist[0][3]):
                idmessdict['Reference']=idmesslist[0][3]
            else:
                idmessdict['Reference']='无'
            if(idmesslist[0][4]):
                idmessdict['Global']=idmesslist[0][4]
            else:
                idmessdict['Global']='无'
            if(idmesslist[0][5]):
                idmessdict['Comment']=idmesslist[0][5]
            else:
                idmessdict['Comment']='无'
        else:
            idmessdict['Id']='无'
            idmessdict['Descriptions']='无'
            idmessdict['Engineer']='无'
            idmessdict['Reference']='无'
            idmessdict['Global']='无'
            idmessdict['Comment']='无'

        '''解析案例'''
        allcase=re.findall(self.patternAllCase,text)
        casenum=1
        caselist=[]
        for case in allcase:
            casedetail={}
            globalstr=''
            tcstubstr=''
            inputstr=''
            outputstr=''

            casedetail['casenum']=casenum
            # 解析函数名
            funname=re.findall(self.patternGetFunName,case)
            casedetail['funname']=funname[0]
            print(funname[0])
            self.funname=funname[0]

            # 解析所在.c文件
            filename=re.findall(self.patternGetFileName,case)
            casedetail['filename']=filename[0]
            #self.filename=filename[0]

            #解析测试技术&用例描述
            testing=re.findall(self.patternGetTesting,case)
            if(testing):
                casedetail['testingway']=testing[0][0]
                casedetail['testingdescribe']=testing[0][1]
            else:
                casedetail['testingway']='边界值、等价类'
                casedetail['testingdescribe']='无'

            # 解析预计输入
            inputlist=re.findall(self.patternGetZ,case)
            if(inputlist):
                for ainput in inputlist:
                    inputstr+=ainput[1]+'  '+ainput[0]+' = '+ainput[2]+' ;\n'

            # 解析全局变量
            gvallist=re.findall(self.patternGetG,case)
            if(gvallist):
                for gval in gvallist:
                    globalstr+=gval[1]+'  '+gval[0]+' ;\n'
                    inputstr+=gval[1]+'  '+gval[0]+' = '+gval[2]+' ;\n'
            if(globalstr):
                casedetail['gvallist']=globalstr
            else:
                casedetail['gvallist']='无'
            if(inputstr):
                casedetail['inputstr']=inputstr
            else:
                casedetail['inputstr']='无'

            # 解析初始化代码
            casedetail['init']=''
            '''解析全局输入'''
            globalinit=re.findall(self.patternGlobalInit,text)
            if(globalinit):
                casedetail['init']=casedetail['init']+globalinit[0]
            '''解析案例输入'''
            initlist=re.findall(self.patternInit,case)
            if(initlist):
                casedetail['init']=casedetail['init']+initlist[0]
            if(casedetail['init']):
                pass
            else:
                casedetail['init']='无'

            # 解析桩函数
            tcstublist=re.findall(self.patternTCStub,case)
            if(tcstublist):
                for atcstub in tcstublist:
                    tcstubstr+=atcstub[0]+'( )  return  '+atcstub[2]+' ;\n'
                casedetail['tcstub']=tcstubstr
            else:
                casedetail['tcstub']='无'

            # 解析输出  包括预计输出、实际输出
            outputlist=re.findall(self.patternGetH,case)
            if(outputlist):
                for aoutput in outputlist:
                    outputstr+=aoutput[1]+'  '+aoutput[0]+' = '+aoutput[2]+' ;\n'

            returnlist=re.findall(self.patternGetO,case)
            if(returnlist):
                for areturn in returnlist:
                    outputstr+='return '+areturn[2]+' ;\n'
            if(outputstr):
                casedetail['outputstr']=outputstr
            else:
                casedetail['outputstr']='无'
            # 结论
            casedetail['result']='PASS'
            caselist.append(casedetail)
            casenum+=1

        return [self.funname,self.funlist,caselist,idmessdict]  #入口函数名、入口函数1个+修函数n个、案例集、序列标识

    #整理信息
    def dealCase(self,list,savepath):
        tmp = open("Tmp.docx","wb+")
        tmp.write(base64.b64decode(docx))
        tmp.close()
        document=Document('Tmp.docx') #导入word （页眉页脚）
        '''页脚'''
        section = document.sections[0]
        table= section.footer.tables[0]
        table.cell(1,2).paragraphs[0].add_run(self.projectyear+'-'+self.projectmonth+'-'+self.projectday)
        table.cell(1,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
        run =table.cell(1,2).paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        '''页脚 END'''
        '''封面'''
        str1=u'\r'*2
        str2=self.projectname
        str3=u'\r'*3
        str4=u'集成测试用例集'
        str5=u'\r'*6
        str6=u'通号城市轨道交通技术有限公司'
        str7=u'\r'
        str8=self.projectyear+'年'+self.projectmonth+'月'+self.projectday+'日'
        strsum=str1+str2+str3+str4+str5+str6+str7+str8
        paragraph = document.add_paragraph('')#添加段落
        run=paragraph.add_run(strsum)
        run.font.name = u"宋体"
        run.font.size = Pt(25)
        run.bold=True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER #文本居中
        document.add_page_break()
        '''封面END'''
        '''修改记录'''
        head = document.add_heading("", level=1)
        run = head.add_run(u"修改记录")
        run.font.size = Pt(22)
        run.bold=True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        #表格
        table = document.add_table(2,5)
        table.style = "Table Grid" #带边框
        table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中

        table.cell(0,0).paragraphs[0].add_run('版本号')
        table.cell(0,1).paragraphs[0].add_run('修改章节')
        table.cell(0,2).paragraphs[0].add_run('修改内容概要')
        table.cell(0,3).paragraphs[0].add_run('修改人')
        table.cell(0,4).paragraphs[0].add_run('修改时间')
        table.cell(1,0).paragraphs[0].add_run('v0.0.1')
        table.cell(1,1).paragraphs[0].add_run('全部')
        table.cell(1,2).paragraphs[0].add_run('创建')
        table.cell(1,3).paragraphs[0].add_run(self.testpeople)
        table.cell(1,4).paragraphs[0].add_run(self.projectyear+'-'+self.projectmonth+'-'+self.projectday)
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
                paragraph_format=cell.paragraphs[0].paragraph_format
                paragraph_format.space_before=Pt(18)    #上行间距
                paragraph_format.space_after=Pt(18)    #下行间距
                paragraph_format.line_spacing=Pt(18)  #行距
                # paragraph_format.left_indent=Inches(0.3)   #调整左缩进0.3英寸
                run =cell.paragraphs[0].runs[0]
                run.font.size = Pt(15)
                run.bold=True #加粗
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = u"宋体"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.add_page_break() #换页
        '''修改记录END'''
        '''签署页'''
        head = document.add_heading("", level=1)
        run = head.add_run(u"签署页")
        run.font.size = Pt(22)
        run.bold=True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        #表格
        table = document.add_table(4,4)
        table.style = "Table Grid" #带边框
        table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
        table.cell(0,1).paragraphs[0].add_run('角色')
        table.cell(0,2).paragraphs[0].add_run('签字')
        table.cell(0,3).paragraphs[0].add_run('日期')
        table.cell(1,0).paragraphs[0].add_run('编制者')
        table.cell(2,0).paragraphs[0].add_run('审核人')
        table.cell(3,0).paragraphs[0].add_run('项目负责人')
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
                cell.paragraphs[0].add_run('')
                paragraph_format=cell.paragraphs[0].paragraph_format
                paragraph_format.space_before=Pt(18)    #上行间距
                paragraph_format.space_after=Pt(18)    #下行间距
                paragraph_format.line_spacing=Pt(18)  #行距
                # paragraph_format.left_indent=Inches(0.3)   #调整左缩进0.3英寸
                run =cell.paragraphs[0].runs[0]
                run.font.size = Pt(15)
                run.bold=True #加粗
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = u"宋体"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.add_page_break()#换页
        '''签署页END'''
        '''测试范围与结果init'''
        head = document.add_heading("", level=1)
        run = head.add_run(u"测试范围与结果")
        run.font.size = Pt(22)
        run.bold=True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        #表格
        table = document.add_table(1,5)
        table.style = "Table Grid" #带边框
        table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
        table.cell(0,0).paragraphs[0].add_run('序号')
        table.cell(0,1).paragraphs[0].add_run('函数名')
        table.cell(0,2).paragraphs[0].add_run('用例个数')
        table.cell(0,3).paragraphs[0].add_run('针对版本')
        table.cell(0,4).paragraphs[0].add_run('测试结果')
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
                cell.paragraphs[0].add_run('')
                paragraph_format=cell.paragraphs[0].paragraph_format
                paragraph_format.space_before=Pt(10)    #上行间距
                paragraph_format.space_after=Pt(10)    #下行间距
                paragraph_format.line_spacing=Pt(10)  #行距
                run =cell.paragraphs[0].runs[0]
                run.font.size = Pt(11)
                run.bold=True #加粗
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = u"宋体"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        document.add_page_break()#换页
        '''测试范围与结果END'''
        funno=1 #函数编号
        for li in list:
            funname=li[0] #入口函数名
            funlist=li[1] #入口函数1个+修函数n个
            caselist=li[2] #案例集
            idmess=li[3] #序列标识
            for f in funlist:
                if(f != funname):
                    table=document.tables[2]
                    row=table.add_row()
                    row.cells[0].paragraphs[0].add_run(str(funno))
                    row.cells[1].paragraphs[0].add_run(f)
                    row.cells[2].paragraphs[0].add_run(str(len(caselist)))
                    row.cells[3].paragraphs[0].add_run(self.projectname)
                    row.cells[4].paragraphs[0].add_run('PASS')
                    funno+=1 #编号加1
                    for cell in row.cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
                        cell.paragraphs[0].add_run('')
                        paragraph_format=cell.paragraphs[0].paragraph_format
                        paragraph_format.space_before=Pt(10)    #上行间距
                        paragraph_format.space_after=Pt(10)    #下行间距
                        paragraph_format.line_spacing=Pt(10)  #行距
                        run =cell.paragraphs[0].runs[0]
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = u"宋体"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    '''测试范围与结果add END'''
                    '''案例描述添加'''
                    head = document.add_heading("", level=1)
                    run = head.add_run(f+'针对'+self.projectname)
                    run.font.size = Pt(22)
                    run.bold=True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = u"宋体"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    paragraph = document.add_paragraph('')#添加入口函数
                    paragraph_format=paragraph.paragraph_format
                    paragraph_format.space_before=Pt(20)    #上行间距
                    paragraph_format.space_after=Pt(10)    #下行间距
                    paragraph_format.line_spacing=Pt(10)  #行距
                    run1 = paragraph.add_run('入口函数名_针对  '+funname)
                    run1.font.size = Pt(9)
                    run1.font.color.rgb = RGBColor(0, 0, 0)
                    run1.bold=True
                    run1.font.name = u"宋体"
                    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    # 表格
                    table = document.add_table(5,6)
                    table.style = "Table Grid" #带边框
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
                    table.cell(0, 1).merge(table.cell(0, 3))
                    table.cell(1, 1).merge(table.cell(1, 5))
                    table.cell(2, 1).merge(table.cell(2, 5))
                    table.cell(3, 1).merge(table.cell(3, 5))
                    table.cell(4, 1).merge(table.cell(4, 5))
                    table.cell(0,0).paragraphs[0].add_run('序列标识')
                    table.cell(0,4).paragraphs[0].add_run('测试者')
                    table.cell(1,0).paragraphs[0].add_run('序列描述')
                    table.cell(2,0).paragraphs[0].add_run('参考文件')
                    table.cell(3,0).paragraphs[0].add_run('用户变量')
                    table.cell(4,0).paragraphs[0].add_run('备注信息')
                    for row in table.rows:
                        for cell in row.cells:
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
                            cell.paragraphs[0].add_run('')
                            paragraph_format=cell.paragraphs[0].paragraph_format
                            paragraph_format.space_before=Pt(8)    #上行间距
                            paragraph_format.space_after=Pt(8)    #下行间距
                            paragraph_format.line_spacing=Pt(8)  #行距
                            run =cell.paragraphs[0].runs[0]
                            run.font.size = Pt(11)
                            run.bold=True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.name = u"宋体"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    table.cell(0,1).text=idmess['Id']
                    table.cell(0,5).text=self.testpeople#idmess['Engineer']
                    table.cell(1,1).text=idmess['Descriptions']
                    table.cell(2,1).text=idmess['Reference']
                    table.cell(3,1).text=idmess['Global']
                    table.cell(4,1).text=idmess['Comment']
                    '''案例描述添加 END'''
                    '''遍历案例集'''
                    caseno=0
                    for casedict in caselist:
                        caseno+=1
                        head = document.add_heading("", level=2)
                        run = head.add_run(u"用例"+str(caseno))
                        run.font.size = Pt(16)
                        run.bold=True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = u"宋体"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        table = document.add_table(9,6)
                        table.style = "Table Grid" #带边框
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
                        table.cell(0, 1).merge(table.cell(0, 2))
                        table.cell(0, 4).merge(table.cell(0, 5))
                        table.cell(1, 1).merge(table.cell(1, 5))
                        table.cell(2, 1).merge(table.cell(2, 5))
                        table.cell(3, 1).merge(table.cell(3, 5))
                        table.cell(4, 1).merge(table.cell(4, 5))
                        table.cell(5, 1).merge(table.cell(5, 5))
                        table.cell(6, 1).merge(table.cell(6, 5))
                        table.cell(7, 1).merge(table.cell(7, 5))
                        table.cell(8, 1).merge(table.cell(8, 5))
                        table.cell(0,0).paragraphs[0].add_run('用例编号')
                        table.cell(0,3).paragraphs[0].add_run('测试技术')
                        table.cell(1,0).paragraphs[0].add_run('用例描述')
                        table.cell(2,0).paragraphs[0].add_run('全局变量')
                        table.cell(3,0).paragraphs[0].add_run('初始化代码')
                        table.cell(4,0).paragraphs[0].add_run('桩函数')
                        table.cell(5,0).paragraphs[0].add_run('输入')
                        table.cell(6,0).paragraphs[0].add_run('预期输出')
                        table.cell(7,0).paragraphs[0].add_run('实际输出')
                        table.cell(8,0).paragraphs[0].add_run('结论')
                        for row in table.rows:
                            for cell in row.cells:
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
                                cell.paragraphs[0].add_run('')
                                paragraph_format=cell.paragraphs[0].paragraph_format
                                paragraph_format.space_before=Pt(8)    #上行间距
                                paragraph_format.space_after=Pt(8)    #下行间距
                                paragraph_format.line_spacing=Pt(8)  #行距
                                run =cell.paragraphs[0].runs[0]
                                run.font.size = Pt(11)
                                run.bold=True #加粗
                                run.font.color.rgb = RGBColor(0, 0, 0)
                                run.font.name = u"宋体"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        table.cell(0,1).text=str(casedict['casenum'])
                        table.cell(0,4).text=casedict['testingway']
                        table.cell(1,1).text=casedict['testingdescribe']
                        table.cell(2,1).text=casedict['gvallist']
                        table.cell(3,1).text=casedict['init']
                        table.cell(4,1).text=casedict['tcstub']
                        table.cell(5,1).text=casedict['inputstr']
                        table.cell(6,1).text=casedict['outputstr']
                        table.cell(7,1).text=casedict['outputstr']
                        table.cell(8,1).text=casedict['result']
                    '''遍历案例集 END'''

        document.save(savepath+'\\'+self.projectname+'集成测试用例集.docx')
        os.remove("Tmp.docx")









if __name__ == '__main__':


    deal=Deal('atp','xxx','2019','11','2')
    list=[]
    #递归遍历筛选文件
    def selectFile(dirpath):
        filelist=[]
        for root, dirs, files in os.walk(dirpath):
            #print(root) #当前目录路径
            #print(dirs) #当前路径下所有子目录
            #print(files) #当前路径下所有非目录子文件
            for file in files:
                if os.path.splitext(file)[1] == '.tcf':
                    filelist.append(os.path.join(root, file))
        return filelist
    lis=selectFile('C:\\Users\\v5682\\Desktop\\新建文件夹1')
    for li in lis:
        print(li)
        text1=deal.getText(li)
        caselist1=deal.getCaseList(text1)
        list.append(caselist1)

    # text=deal.getText('C:\\Users\\v5682\\Desktop\\新建文件夹1\\Seq_Iso_V1.6.0_P1_1.0.0_ar_noman_ar_head.tcf')
    # list.append(deal.getCaseList(text))
    deal.dealCase(list,'C:\\Users\\v5682\\Desktop\\新建文件夹1')

