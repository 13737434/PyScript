from docx import *
from docx.enum.style import WD_STYLE_TYPE
document = Document()
document.add_heading(u'MS WORD写入测试',0)
document.add_heading(u'一级标题',1)
document.add_heading(u'二级标题',2)
# sections = document.sections
# current_section = document.sections[-1]
# print(current_section.start_type)

# 保存新创建的 word 文档
document.save('testDoc.docx')