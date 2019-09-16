from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import os

# 创建一个已存在的 word 文档的对象
document = Document()
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')#添加段落
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
paragraph.style = 'List Bullet'
paragraph1 = document.add_paragraph('Lorem ipsum do.')
paragraph2 = document.add_paragraph('\n\n')
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')#在其上方插入一个新段落

paragraph = document.add_paragraph('Lorem ipsum ')
run = paragraph.add_run('dolor')#.add_run()段落上的方法添加更多内容
run.bold = True#Run对象具有a .bold和.italic属性
paragraph.add_run(' sit amet.')

paragraph = document.add_paragraph('Normal text, ')#应用字符样式
run = paragraph.add_run('text with emphasis.')
run.style = 'Emphasis'

document.add_page_break()#添加分页

table = document.add_table(rows=2, cols=2)#添加表
cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'

#table.style = 'LightShading-Accent1'

# 保存新创建的 word 文档
document.save('testDoc.docx')