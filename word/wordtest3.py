from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT,WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

document=Document()
'''页眉页脚'''
section = document.sections[0]
section.header_distance=Inches(0)#表示从页面上边缘到标题顶边的距离的对象
header = section.header
table=header.add_table(2,12,Inches(17))
table.cell(0, 0).merge(table.cell(1, 4))
table.cell(0, 5).merge(table.cell(0, 6))
table.cell(0, 7).merge(table.cell(0, 11))
table.cell(1, 5).merge(table.cell(1, 6))
table.cell(1, 7).merge(table.cell(1, 11))
table.cell(0,5).paragraphs[0].add_run('项目名称')
table.cell(1,5).paragraphs[0].add_run('文件名称')
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
        cell.paragraphs[0].add_run('')
        paragraph_format=cell.paragraphs[0].paragraph_format
        paragraph_format.space_before=Pt(10)    #上行间距
        paragraph_format.space_after=Pt(10)    #下行间距
        paragraph_format.line_spacing=Pt(10)  #行距
        run =cell.paragraphs[0].runs[0]
        run.font.size = Pt(11)
        run.bold=True #加粗
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
header.add_paragraph('').add_run('\t'*12).underline = True #添加横线






'''封面'''
str1=u'\r'*3
str2=u'ATP_V1.6.0_P1_1.0.0'
str3=u'\r'*3
str4=u'ar_hd_adap模块单元测试用例集'
str5=u'\r'*7
str6=u'通号城市轨道交通技术有限公司'
str7=u'\r'
str8=u'2019年08月27日'
strsum=str1+str2+str3+str4+str5+str6+str7+str8

paragraph = document.add_paragraph('')#添加段落
run=paragraph.add_run(strsum)
run.font.name = u"宋体"
run.font.size = Pt(25)
run.bold=True
run.font.color.rgb = RGBColor(0, 0, 0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER #文本居中

document.add_page_break()
'''封面END'''

'''修改记录'''
head = document.add_heading("", level=1)
run = head.add_run(u"修改记录")
run.font.size = Pt(22)
run.bold=True
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.name = u"宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

#表格
table = document.add_table(2,5)
table.style = "Table Grid" #带边框
table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中

table.cell(0,0).paragraphs[0].add_run('版本号')
table.cell(0,1).paragraphs[0].add_run('修改章节')
table.cell(0,2).paragraphs[0].add_run('修改内容概要')
table.cell(0,3).paragraphs[0].add_run('修改人')
table.cell(0,4).paragraphs[0].add_run('修改时间')
table.cell(1,0).paragraphs[0].add_run('v0.0.1')
table.cell(1,1).paragraphs[0].add_run('全部')
table.cell(1,2).paragraphs[0].add_run('创建')
table.cell(1,3).paragraphs[0].add_run('xxx')
table.cell(1,4).paragraphs[0].add_run('20190827')


for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
        paragraph_format=cell.paragraphs[0].paragraph_format
        paragraph_format.space_before=Pt(18)    #上行间距
        paragraph_format.space_after=Pt(18)    #下行间距
        paragraph_format.line_spacing=Pt(18)  #行距
        # paragraph_format.left_indent=Inches(0.3)   #调整左缩进0.3英寸
        run =cell.paragraphs[0].runs[0]
        run.font.size = Pt(15)
        run.bold=True #加粗
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.add_page_break() #换页
'''修改记录END'''

'''签署页'''
head = document.add_heading("", level=1)
run = head.add_run(u"签署页")
run.font.size = Pt(22)
run.bold=True
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.name = u"宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#表格
table = document.add_table(4,4)
table.style = "Table Grid" #带边框
table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
table.cell(0,1).paragraphs[0].add_run('角色')
table.cell(0,2).paragraphs[0].add_run('签字')
table.cell(0,3).paragraphs[0].add_run('日期')
table.cell(1,0).paragraphs[0].add_run('编制者')
table.cell(2,0).paragraphs[0].add_run('审核人')
table.cell(3,0).paragraphs[0].add_run('项目负责人')
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
        cell.paragraphs[0].add_run('')
        paragraph_format=cell.paragraphs[0].paragraph_format
        paragraph_format.space_before=Pt(18)    #上行间距
        paragraph_format.space_after=Pt(18)    #下行间距
        paragraph_format.line_spacing=Pt(18)  #行距
        # paragraph_format.left_indent=Inches(0.3)   #调整左缩进0.3英寸
        run =cell.paragraphs[0].runs[0]
        run.font.size = Pt(15)
        run.bold=True #加粗
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
document.add_page_break()#换页
'''签署页END'''

'''测试范围与结果'''
head = document.add_heading("", level=1)
run = head.add_run(u"测试范围与结果")
run.font.size = Pt(22)
run.bold=True
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.name = u"宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
#表格
table = document.add_table(2,5)
table.style = "Table Grid" #带边框
table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
table.cell(0,0).paragraphs[0].add_run('序号')
table.cell(0,1).paragraphs[0].add_run('函数名')
table.cell(0,2).paragraphs[0].add_run('用例个数')
table.cell(0,3).paragraphs[0].add_run('针对版本')
table.cell(0,4).paragraphs[0].add_run('测试结果')
table.cell(1,0).paragraphs[0].add_run('1')
table.cell(1,1).paragraphs[0].add_run('appl_ar_hd_adap_init')
table.cell(1,2).paragraphs[0].add_run('1')
table.cell(1,3).paragraphs[0].add_run('v1.6.0_p1_1.0.0')
table.cell(1,4).paragraphs[0].add_run('pass')
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
        cell.paragraphs[0].add_run('')
        paragraph_format=cell.paragraphs[0].paragraph_format
        paragraph_format.space_before=Pt(10)    #上行间距
        paragraph_format.space_after=Pt(10)    #下行间距
        paragraph_format.line_spacing=Pt(10)  #行距
        run =cell.paragraphs[0].runs[0]
        run.font.size = Pt(11)
        run.bold=True #加粗
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

document.add_page_break()#换页
'''测试范围与结果END'''

'''案例'''
head = document.add_heading("", level=1)
run = head.add_run(u"appl_ar_hd_adap_init_针对V1.6.0_P1_1.0.0")
run.font.size = Pt(22)
run.bold=True
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.name = u"宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# 表格
table = document.add_table(5,6)
table.style = "Table Grid" #带边框
table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
table.cell(0, 1).merge(table.cell(0, 3))
table.cell(1, 1).merge(table.cell(1, 5))
table.cell(2, 1).merge(table.cell(2, 5))
table.cell(3, 1).merge(table.cell(3, 5))
table.cell(4, 1).merge(table.cell(4, 5))
table.cell(0,0).paragraphs[0].add_run('序列标识')
table.cell(0,4).paragraphs[0].add_run('测试者')
table.cell(1,0).paragraphs[0].add_run('序列描述')
table.cell(2,0).paragraphs[0].add_run('参考文件')
table.cell(3,0).paragraphs[0].add_run('用户变量')
table.cell(4,0).paragraphs[0].add_run('备注信息')
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
        cell.paragraphs[0].add_run('')
        paragraph_format=cell.paragraphs[0].paragraph_format
        paragraph_format.space_before=Pt(8)    #上行间距
        paragraph_format.space_after=Pt(8)    #下行间距
        paragraph_format.line_spacing=Pt(8)  #行距
        run =cell.paragraphs[0].runs[0]
        run.font.size = Pt(11)
        run.bold=True #加粗
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 案例x
head = document.add_heading("", level=2)
run = head.add_run(u"案例1")
run.font.size = Pt(11)
run.bold=True
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.name = u"宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# 表格
table = document.add_table(9,6)
table.style = "Table Grid" #带边框
table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
table.cell(0, 1).merge(table.cell(0, 2))
table.cell(0, 4).merge(table.cell(0, 5))
table.cell(1, 1).merge(table.cell(1, 5))
table.cell(2, 1).merge(table.cell(2, 5))
table.cell(3, 1).merge(table.cell(3, 5))
table.cell(4, 1).merge(table.cell(4, 5))
table.cell(5, 1).merge(table.cell(5, 5))
table.cell(6, 1).merge(table.cell(6, 5))
table.cell(7, 1).merge(table.cell(7, 5))
table.cell(8, 1).merge(table.cell(8, 5))
table.cell(0,0).paragraphs[0].add_run('用例编号')
table.cell(0,3).paragraphs[0].add_run('测试技术')
table.cell(1,0).paragraphs[0].add_run('用例描述')
table.cell(2,0).paragraphs[0].add_run('全局变量')
table.cell(3,0).paragraphs[0].add_run('初始化代码')
table.cell(4,0).paragraphs[0].add_run('桩函数')
table.cell(5,0).paragraphs[0].add_run('输入')
table.cell(6,0).paragraphs[0].add_run('预期输出')
table.cell(7,0).paragraphs[0].add_run('实际输出')
table.cell(8,0).paragraphs[0].add_run('结论')
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #竖直居中
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #水平居中
        cell.paragraphs[0].add_run('')
        paragraph_format=cell.paragraphs[0].paragraph_format
        paragraph_format.space_before=Pt(8)    #上行间距
        paragraph_format.space_after=Pt(8)    #下行间距
        paragraph_format.line_spacing=Pt(8)  #行距
        run =cell.paragraphs[0].runs[0]
        run.font.size = Pt(11)
        run.bold=True #加粗
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

'''案例END'''









document.save('0000.docx')
