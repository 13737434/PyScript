import openpyxl
from openpyxl import load_workbook
from docx import *
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT,WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
wb = load_workbook('test.xlsx')
ws=wb.active
list=[]
for i in range(1,10): #行号
    tmp=[]
    for nn in ['A','B','C','D','E','F','G','H','I','J','K','L','M','N','O']: #列号
        tmp.append(ws[nn+str(i)])
    list.append(tmp)

document=Document('Tmp.docx')
for t in list:

    table = document.add_table(2,5)
    table.style = "Table Grid" #带边框
    table.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中
    row=table.add_row()
    row.cells[0].paragraphs[0].add_run('str(funno)')
    row.cells[1].paragraphs[0].add_run('funname')
    row.cells[2].paragraphs[0].add_run('str(len(caselist))')
    row.cells[3].paragraphs[0].add_run('self.projectname')
    row.cells[4].paragraphs[0].add_run('PASS')